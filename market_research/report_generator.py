import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
import re
from datetime import datetime
from urllib.parse import urlparse


def extract_company_name(content):
    """Extract company name from the report content"""
    # Try to find company name in the first few lines
    first_lines = content.split('\n')[:10]
    for line in first_lines:
        # Look for company name in headers or first paragraph
        if any(word in line.lower() for word in ['overview', 'analysis', 'report']):
            # Extract potential company name
            words = line.strip('#').strip().split()
            for i, word in enumerate(words):
                if word.lower() in ['of', 'for', 'about']:
                    return ' '.join(words[i+1:i+3])  # Take next 2 words after 'of/for/about'
    
    # Fallback to first meaningful words
    for line in first_lines:
        if line.strip() and not line.startswith('#'):
            words = line.strip().split()
            return ' '.join(words[:2])  # Take first two words
    
    return "Company"  # Default fallback

def sanitize_filename(filename):
    """Sanitize the filename to be safe for all operating systems"""
    # Remove or replace invalid characters
    filename = re.sub(r'[<>:"/\\|?*]', '', filename)
    # Remove multiple spaces
    filename = ' '.join(filename.split())
    return filename.strip()

def setup_document_styles(doc):
    """Set up document styles for consistent formatting"""
    # Normal text style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # Heading styles
    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = 'Calibri'
        heading_style.font.bold = True
        heading_style.font.color.rgb = RGBColor(21, 101, 192)  # Blue color for headings
        heading_style.paragraph_format.space_before = Pt(16)
        heading_style.paragraph_format.space_after = Pt(8)
        heading_style.paragraph_format.keep_with_next = True

    # Quote style - check if it exists first
    try:
        quote_style = doc.styles.add_style('Quote', WD_STYLE_TYPE.PARAGRAPH)
        quote_style.base_style = doc.styles['Normal']
    except ValueError:  # Style already exists
        quote_style = doc.styles['Quote']
    
    quote_style.font.italic = True
    quote_style.font.color.rgb = RGBColor(66, 66, 66)
    quote_style.paragraph_format.left_indent = Inches(0.5)
    quote_style.paragraph_format.space_before = Pt(12)
    quote_style.paragraph_format.space_after = Pt(12)

def extract_markdown_links(content):
    """Extract markdown-style links with their context"""
    links = []
    # Match markdown links [text](url)
    pattern = r'\[([^\]]+)\]\(([^)]+)\)'
    
    lines = content.split('\n')
    current_section = "General"
    
    for i, line in enumerate(lines):
        if line.strip().startswith('#'):
            current_section = line.strip('#').strip()
        
        matches = re.finditer(pattern, line)
        for match in matches:
            text, url = match.groups()
            context = line.strip()
            
            # Categorize the source
            category = categorize_source(text, url)
            
            links.append({
                'text': text,
                'url': url,
                'context': context,
                'section': current_section,
                'category': category
            })
    
    return links

def categorize_source(text, url):
    """Categorize the source based on text and URL"""
    text_lower = text.lower()
    domain = urlparse(url).netloc.lower()
    
    categories = {
        'Market Research': ['research', 'market analysis', 'forecast', 'report', 'grandview', 'frost', 'marketsandmarkets'],
        'Scientific Publications': ['nature', 'science', 'cell', 'lancet', 'journal', 'pubmed', 'nih.gov', 'research paper'],
        'Regulatory & Government': ['fda', 'ema', 'gov', 'regulation', 'guidance', 'guidelines'],
        'Industry News': ['news', 'press', 'article', 'forbes', 'reuters', 'bloomberg'],
        'Company Resources': ['company', 'corporate', 'investor', 'presentation', 'annual report'],
        'Healthcare Organizations': ['hospital', 'clinic', 'medical center', 'health', 'care', 'who.int']
    }
    
    for category, keywords in categories.items():
        if any(keyword in text_lower or keyword in domain for keyword in keywords):
            return category
    
    return "Other Sources"

def format_sources_section(doc, links):
    """Format sources section with categorized links"""
    if not links:
        return
    
    doc.add_heading('Sources', 1)
    
    # Group links by category
    categories = {}
    for link in links:
        category = link['category']
        if category not in categories:
            categories[category] = []
        categories[category].append(link)
    
    # Add sources by category
    for category, category_links in categories.items():
        doc.add_heading(category, 2)
        for link in category_links:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"{link['text']}").bold = True
            p.add_run(f" (Accessed: {get_current_date()})")
            p.add_run(f"\nURL: {link['url']}")
            if link['context'] != link['text']:
                p.add_run(f"\nContext: {link['context']}")

def get_current_date():
    """Get current date in a readable format"""
    return datetime.now().strftime("%B %d, %Y")

def process_inline_citations(content):
    """Process inline citations to use markdown-style links"""
    # Replace raw URL citations with markdown-style links
    pattern = r'\(source: (http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+)\)'
    return re.sub(pattern, r'([source](\1))', content)

def generate_report_file(research_data):
    """Generate a formatted Word document from the research data"""
    doc = Document()
    setup_document_styles(doc)
    
    # Convert CrewAI output to string if it's not already
    content = str(research_data)
    
    # Process inline citations
    content = process_inline_citations(content)
    
    # Extract markdown links before processing content
    links = extract_markdown_links(content)
    
    # Title
    title = doc.add_heading('Market Research Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add generation date
    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_paragraph.add_run(f"Generated on {get_current_date()}")
    
    # Process content sections
    sections = []
    current_section = []
    
    for line in content.split('\n'):
        if line.strip().startswith('#'):
            if current_section:
                sections.append('\n'.join(current_section))
                current_section = []
        current_section.append(line.strip())
    
    if current_section:
        sections.append('\n'.join(current_section))
    
    # Process each section
    for section in sections:
        if section.strip():
            lines = section.split('\n')
            for line in lines:
                if line.strip().startswith('#'):
                    level = min(line.count('#'), 9)
                    text = line.replace('#', '').strip()
                    doc.add_heading(text, level)
                elif line.strip().startswith('>'):
                    text = line.strip('> ').strip()
                    p = doc.add_paragraph(text, style='Quote')
                elif line.strip().startswith('- ') or line.strip().startswith('* '):
                    text = line.strip('- ').strip('* ').strip()
                    p = doc.add_paragraph(text)
                    p.style = 'List Bullet'
                elif line.strip().startswith('1. '):
                    text = line.strip('1. ').strip()
                    p = doc.add_paragraph(text)
                    p.style = 'List Number'
                else:
                    if line.strip():
                        doc.add_paragraph(line.strip())
    
    # Add sources section with categorized links
    format_sources_section(doc, links)
    
    # Save the document
    if not os.path.exists('reports'):
        os.makedirs('reports')
    
    company_name = extract_company_name(content)
    filename = sanitize_filename(f"Market Analysis of {company_name}.docx")
    report_path = os.path.join('reports', filename)
    
    doc.save(report_path)
    
    return report_path, content

def generate_report_content(research_data):
    """Generate comprehensive report content from research data"""
    content = []
    
    # Title
    content.append(f"# Comprehensive Market Research Report: {research_data['title']}\n")
    
    # Executive Summary
    content.append("## Executive Summary")
    content.append("### Key Findings")
    content.append(research_data.get('key_findings', ''))
    content.append("\n### Market Highlights")
    content.append(research_data.get('market_highlights', ''))
    content.append("\n### Strategic Recommendations")
    content.append(research_data.get('strategic_recommendations', ''))
    content.append("\n")
    
    # Company Overview
    content.append("## Company Overview")
    content.append("### Company History and Background")
    content.append(research_data.get('company_history', ''))
    content.append("\n### Mission and Vision")
    content.append(research_data.get('mission_vision', ''))
    content.append("\n### Key Leadership and Organizational Structure")
    content.append(research_data.get('leadership', ''))
    content.append("\n")
    
    # Product/Service Analysis
    content.append("## Product/Service Analysis")
    content.append("### Core Offerings")
    content.append(research_data.get('core_offerings', ''))
    content.append("\n### Key Features and Benefits")
    content.append(research_data.get('key_features', ''))
    content.append("\n### Technology and Innovation")
    content.append(research_data.get('technology', ''))
    content.append("\n### Product Development Pipeline")
    content.append(research_data.get('development_pipeline', ''))
    content.append("\n")
    
    # Market Analysis
    content.append("## Market Analysis")
    content.append("### Industry Overview and Size")
    content.append(research_data.get('industry_overview', ''))
    content.append("\n### Market Trends and Dynamics")
    content.append(research_data.get('market_trends', ''))
    content.append("\n### Growth Drivers and Barriers")
    content.append(research_data.get('growth_drivers', ''))
    content.append("\n### Competitive Landscape")
    content.append(research_data.get('competitive_landscape', ''))
    content.append("\n")
    
    # Business Strategy
    content.append("## Business Strategy")
    content.append("### Go-to-Market Strategy")
    content.append(research_data.get('go_to_market', ''))
    content.append("\n### Revenue Model")
    content.append(research_data.get('revenue_model', ''))
    content.append("\n### Strategic Partnerships")
    content.append(research_data.get('partnerships', ''))
    content.append("\n### Geographic Presence")
    content.append(research_data.get('geographic_presence', ''))
    content.append("\n")
    
    # Financial Analysis
    content.append("## Financial Analysis")
    content.append("### Revenue and Growth Metrics")
    content.append(research_data.get('revenue_metrics', ''))
    content.append("\n### Funding and Investments")
    content.append(research_data.get('funding', ''))
    content.append("\n### Key Financial Indicators")
    content.append(research_data.get('financial_indicators', ''))
    content.append("\n### Future Projections")
    content.append(research_data.get('projections', ''))
    content.append("\n")
    
    # SWOT Analysis
    content.append("## SWOT Analysis")
    content.append("### Strengths")
    content.append(research_data.get('strengths', ''))
    content.append("\n### Weaknesses")
    content.append(research_data.get('weaknesses', ''))
    content.append("\n### Opportunities")
    content.append(research_data.get('opportunities', ''))
    content.append("\n### Threats")
    content.append(research_data.get('threats', ''))
    content.append("\n")
    
    # Future Outlook
    content.append("## Future Outlook")
    content.append("### Growth Opportunities")
    content.append(research_data.get('growth_opportunities', ''))
    content.append("\n### Potential Challenges")
    content.append(research_data.get('challenges', ''))
    content.append("\n### Industry Predictions")
    content.append(research_data.get('predictions', ''))
    content.append("\n### Strategic Recommendations")
    content.append(research_data.get('future_recommendations', ''))
    content.append("\n")
    
    # Conclusion
    content.append("## Conclusion")
    content.append("### Key Takeaways")
    content.append(research_data.get('key_takeaways', ''))
    content.append("\n### Strategic Implications")
    content.append(research_data.get('implications', ''))
    content.append("\n### Final Recommendations")
    content.append(research_data.get('final_recommendations', ''))
    content.append("\n")
    
    # Sources
    content.append("## Sources")
    if research_data.get('sources'):
        for source in research_data['sources']:
            content.append(f"[{source['title']}]({source['url']})")
    
    return '\n'.join(content) 